from docx import Document
import pandas as pd 
import matplotlib.pyplot as plt

students=pd.read_excel('data.xlsx')
students.sort_values(by='Score',inplace=True,ascending=False)

plt.bar(students.Name,students.Score,color='orange')

plt.title('Students score',fontsize=16)
plt.xlabel('Name')
plt.ylabel('score')

plt.xticks(students.Name,rotation='90')
plt.tight_layout()
# plt.show()
imgname='data.jpg'
plt.savefig(imgname)

document=Document()

document.add_heading('数据分析报告',level=0)


first_student=students.iloc[0,:]['Name']
first_score=students.iloc[0,:]['Score']

p=document.add_paragraph('分数排名第一的学生是')
p.add_run(str(first_student)).bold=True
p.add_run(',他的分数是')
p.add_run(str(first_score)).bold=True

p1=document.add_paragraph(f'总共有{len(students.Name)}名学生参加xxx,xxx:')

table=document.add_table(rows=len(students.Name)+1,cols=2)

table.style='LightShading-Accent1'

table.cell(0,0).text='学生名称'
table.cell(0,1).text='学生分数'

for i,(index,row) in enumerate(students.iterrows()):
    table.cell(i+1,0).text=str(row['Name'])
    table.cell(i+1,1).text=str(row['Score'])

document.add_picture(imgname)
document.save('s.docx')
print('done')
